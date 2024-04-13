from docx import Document
from docx.shared import Pt
from core.image_processing import process_image, filter_dataframe
from core.utilities import set_section_parameters, pixels_to_inches, add_paragraph, scale_to_A4


def generate_output_document(images, output_name, scale_font=1, scale_space=0.8, scale_fields=0.37, lang="rus", detection_threshold=25):
    """
    Генерирует документ Word на основе изображений с текстом.

    :param images: Список изображений в формате PIL.Image.Image.
    :param output_name: Имя файла для сохранения результата.
    :param scale_font: Масштабирование размера шрифта (по умолчанию 1).
    :param scale_space: Масштабирование пространства между абзацами (по умолчанию 0.8).
    :param scale_fields: Масштабирование полей (по умолчанию 0.37).
    :param lang: Язык распознавания текста (по умолчанию "rus").
    :param detection_threshold: Пороговое значение уверенности распознавания текста (по умолчанию 25).
    """
    doc = Document()
    section = doc.sections[0]

    for image_index, image in enumerate(images):
        df, fields, dpi = process_image(image, horizontal_shift_threshold=50, lang=lang)
        scale_width, scale_height = scale_to_A4( df["width"][0], df["height"][0], dpi)
        section = set_section_parameters(section, df["width"][0]*scale_width, df["height"][0]*scale_height, dpi)
        df_filtered = filter_dataframe(df, threshold=detection_threshold)
        top_before = 0
        for _, row in df_filtered.iterrows():
            left = int(pixels_to_inches(row['left'], dpi)*scale_width)
            top = int(pixels_to_inches(row['top'], dpi)*scale_height)
            width = int(pixels_to_inches(row['width'], dpi)*scale_width)
            height = int(pixels_to_inches(row['height'], dpi)*scale_height)
            space_adjustment = int(abs(top - top_before) * scale_space)

            if all(element == "!____!" for element in row['text']):
                add_paragraph(doc, "_"*int((width/(Pt(12) * scale_fields))), left, space_adjustment, Pt(12))
            else:
                id_field = 0
                text = []
                for word in row['text']:
                    if word == "!____!":
                        field_id_value = row['field_id'][id_field]
                        field_width = int(pixels_to_inches(fields.loc[fields['field_id'] == field_id_value, 'width'].values[0], dpi) * scale_width)
                        text.append("_"*int((field_width/(height * scale_font * scale_fields))))
                        id_field += 1
                    else:
                        text.append(word)
                add_paragraph(doc, " ".join(text), left, space_adjustment, int(height * scale_font))

            top_before = top + height

        if image_index < len(images) - 1:
            doc.add_page_break()
            section = doc.add_section()

    doc.save(output_name)

def generate_output_text(images, lang="rus+eng", detection_threshold=25) -> str:
    """
    Генерирует текстовую строку на основе изображений с текстом.

    :param images: Список изображений в формате PIL.Image.Image.
    :param lang: Язык распознавания текста (по умолчанию "rus").
    :param detection_threshold: Пороговое значение уверенности распознавания текста (по умолчанию 25).
    :return: Текстовая строка.
    """
    output_text = ""

    for image_index, image in enumerate(images):
        df, _, dpi = process_image(image, horizontal_shift_threshold=50, lang=lang)
        df_filtered = filter_dataframe(df, threshold=detection_threshold)
        found_start = False
        for _, row in df_filtered.iterrows():
            if found_start:
                text = " ".join(row['text']).lower()
                output_text += text + "\n"
            else:
                text = " ".join(row['text']).lower()
                output_text += text + "\n"
                found_start = True

    return output_text